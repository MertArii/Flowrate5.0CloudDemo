"""Doküman -> parçalama (chunking) -> embedding -> pgvector."""
from __future__ import annotations

import re

from pypdf import PdfReader
from app.config import settings
from app.rag import ollama_client, store

# T-code / kod-listesi tarzı girişleri tespit eder: 2 harfle başlayıp içinde
# en az bir rakam geçen kısa kodlar (ME28, MM60, MB5T, ME52N...). Bu, sözlük/
# referans tarzı dokümanlarda (SAP T-code listesi gibi) her girişi kendi
# başına bir chunk yapmayı sağlar — karakter-bazlı bölme bir girişi ortadan
# kesebilir veya birden fazla ilgisiz girişi tek chunk'ta birleştirebilir.
_ENTRY_PATTERN = re.compile(r"\b[A-Z]{2}[A-Z0-9]*\d[A-Z0-9]*\b")
_MIN_ENTRIES = 3  # bundan az eşleşme varsa karakter-bazlı yönteme düş


def chunk_text(text: str, size: int, overlap: int) -> list[str]:
    """Basit karakter-tabanlı, örtüşmeli parçalama."""
    text = " ".join(text.split())
    chunks, start = [], 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start = end - overlap
    return [c for c in chunks if c.strip()]


def chunk_by_entries(text: str) -> list[str] | None:
    """Kod-listesi tarzı içerikte her girişi (kod + açıklaması) ayrı bir
    chunk yapar. Yeterli sayıda giriş bulunamazsa None döner — çağıran taraf
    karakter-bazlı yönteme düşmeli."""
    text = " ".join(text.split())
    matches = list(_ENTRY_PATTERN.finditer(text))
    if len(matches) < _MIN_ENTRIES:
        return None

    chunks = []
    if matches[0].start() > 0:
        onsoz = text[: matches[0].start()].strip()
        if onsoz:
            chunks.append(onsoz)
    for i, m in enumerate(matches):
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        parca = text[m.start():end].strip()
        if parca:
            chunks.append(parca)
    return chunks


def read_file(path: str) -> str:
    """Dosya tipini UZANTIYA değil, gerçek içeriğine (baştaki imza baytları)
    bakarak tespit eder — dosya adı uzantısız/yanlış gelse bile çalışır.
    PDF: '%PDF' imzası. Word (.docx): ZIP imzası ('PK') + python-docx ile
    doğru ayrıştırma (docx aslında bir ZIP arşividir — ZIP'i düz metin gibi
    okumaya çalışmak anlamsız/bozuk içerik ve NUL bayt hatalarına yol açar,
    bu daha önce tam olarak bu şekilde patlamıştı).
    Metin dosyalarında UTF-8 başarısız olursa yaygın Türkçe kodlamaları
    (Windows-1254, ISO-8859-9) dener, o da olmazsa latin-1'e düşer (asla
    hata vermez, en kötü ihtimalle bazı karakterler bozuk görünür)."""
    with open(path, "rb") as f:
        head = f.read(8)

    if head.startswith(b"%PDF"):
        reader = PdfReader(path)
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        return _sanitize(text)

    if head[:2] == b"PK":
        return _sanitize(_read_docx(path))

    with open(path, "rb") as f:
        raw = f.read()
    for enc in ("utf-8", "cp1254", "iso-8859-9"):
        try:
            return _sanitize(raw.decode(enc))
        except UnicodeDecodeError:
            continue
    return _sanitize(raw.decode("latin-1"))


def _read_docx(path: str) -> str:
    from docx import Document
    try:
        doc = Document(path)
    except Exception as e:
        # ZIP imzalı ama .docx değil (xlsx/pptx/jar olabilir) — anlamsız
        # ikili veriyi metin gibi kaydetmek yerine açıkça hata ver.
        raise ValueError(
            f"ZIP tabanlı dosya ama Word (.docx) olarak ayrıştırılamadı "
            f"(xlsx/pptx olabilir mi?): {e}"
        )
    parcalar = [p.text for p in doc.paragraphs if p.text.strip()]
    for tablo in doc.tables:
        for satir in tablo.rows:
            parcalar.append(" | ".join(hucre.text for hucre in satir.cells))
    return "\n".join(parcalar)


def _sanitize(text: str) -> str:
    """Postgres'in TEXT sütunlarının kabul etmediği NUL (0x00) baytlarını
    temizler. Bazı karmaşık/gömülü fontlu PDF'lerde pypdf bunları üretebiliyor."""
    return text.replace("\x00", "")


async def ingest_file(path: str, source: str, title: str) -> int:
    """Dokümanı parçalayıp attachment_vectors'e (RAG Katman 2) yazar.
    Eklenen parça sayısını döner.

    Önce giriş-bazlı bölmeyi dener (kod-listesi tarzı içerik için); yeterli
    giriş bulunamazsa karakter-bazlı yönteme düşer (prose/anlatı metinler)."""
    raw = read_file(path)
    pieces = chunk_by_entries(raw)
    if pieces is None:
        pieces = chunk_text(raw, settings.chunk_size, settings.chunk_overlap)
    embedded = [(p, await ollama_client.embed(p)) for p in pieces]
    return store.add_knowledge_chunks(source or title, embedded)
